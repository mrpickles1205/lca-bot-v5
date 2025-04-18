
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import datetime
import random

def generate_lci_data():
    return pd.DataFrame({
        'Life Cycle Stage': ['Materials', 'Manufacturing', 'Use Phase', 'End-of-Life'],
        'Energy Use (MJ)': [random.uniform(80, 120), random.uniform(50, 100), random.uniform(10, 20), random.uniform(15, 30)],
        'GHG Emissions (kg CO2-eq)': [random.uniform(5, 10), random.uniform(8, 12), random.uniform(1, 3), random.uniform(2, 4)],
        'Water Use (L)': [random.uniform(20, 40), random.uniform(10, 30), random.uniform(1, 5), random.uniform(5, 15)]
    })

def create_charts(df):
    chart_paths = []
    for column in df.columns[1:]:
        fig, ax = plt.subplots()
        ax.bar(df['Life Cycle Stage'], df[column])
        ax.set_title(column)
        chart_path = f"{column.replace(' ', '_')}.png"
        fig.savefig(chart_path)
        chart_paths.append(chart_path)
        plt.close(fig)
    return chart_paths

def generate_docx(product_name, df, chart_paths):
    doc = Document()
    doc.add_heading(f'ISO LCA Report: {product_name}', 0)
    doc.add_paragraph(f'Date: {datetime.date.today()}')

    doc.add_heading("Life Cycle Inventory (LCI)", level=1)
    doc.add_paragraph("This section details the inventory of energy, greenhouse gas emissions, and water use across each life cycle phase.")

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(round(val, 2)) if isinstance(val, (int, float)) else str(val)

    doc.add_page_break()
    doc.add_heading("Visualizations", level=1)
    for chart in chart_paths:
        doc.add_paragraph(chart.replace('_', ' ').replace('.png', ''))
        doc.add_picture(chart, width=Inches(5.5))

    filename = f"LCA_Report_{product_name.replace(' ', '_')}.docx"
    doc.save(filename)
    return filename

st.title("🌿 LCA Bot with Visual Report")
product = st.text_input("Enter the product name:", "Electric Toothbrush")

if st.button("Generate LCA Report"):
    st.info("Generating simulated inventory and visualizations...")
    df = generate_lci_data()
    chart_paths = create_charts(df)
    report_path = generate_docx(product, df, chart_paths)

    with open(report_path, "rb") as file:
        st.download_button("📄 Download Word Report", file, file_name=report_path)
